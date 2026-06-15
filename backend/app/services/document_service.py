import os
from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument
import chardet
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.services.vectorstore_service import vector_store_service
from app.models.database import Document
from app.config import get_settings

settings = get_settings()


def parse_file(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    if ext == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif ext == ".docx":
        doc = DocxDocument(BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        # try to detect encoding
        detection = chardet.detect(file_bytes)
        encoding = detection.get("encoding", "utf-8")
        text = file_bytes.decode(encoding or "utf-8", errors="replace")
    return text


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )
    return splitter.split_text(text)


async def process_document(file_bytes: bytes, filename: str, file_type: str, db_session) -> Document:
    # 1. Parse
    text = parse_file(file_bytes, filename)

    # 2. Chunk
    chunks = chunk_text(text, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)

    # 3. Create DB record
    db_doc = Document(
        filename=filename,
        file_type=file_type,
        file_size=len(file_bytes),
        chunk_count=len(chunks),
        status="indexing"
    )
    db_session.add(db_doc)
    await db_session.commit()
    await db_session.refresh(db_doc)

    # 4. Store in FAISS
    metadatas = [{"doc_id": db_doc.id, "chunk_index": i, "source": filename} for i in range(len(chunks))]
    vector_store_service.add_documents(chunks, metadatas, db_doc.id)

    # 5. Update DB record
    db_doc.status = "indexed"
    await db_session.commit()
    await db_session.refresh(db_doc)

    return db_doc
